"""
DocuChat - A Simple Document Q&A System
Tutorial: FastAPI + LangChain + OpenAI + ChromaDB
"""

import os
import tempfile
from typing import List, Optional
from contextlib import asynccontextmanager
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse
from pydantic import BaseModel
import chromadb
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.docstore.document import Document
import PyPDF2
import docx

# Pydantic models for request/response
class ChatRequest(BaseModel):
    question: str

class ChatResponse(BaseModel):
    answer: str
    sources: Optional[List[str]] = None

# Global variables for our AI components
vectorstore = None
qa_chain = None
embeddings = None
llm = None

# Initialize ChromaDB client with settings to disable telemetry
chroma_client = chromadb.Client(
    settings=chromadb.config.Settings(
        anonymized_telemetry=False,
        allow_reset=True
    )
)

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize and cleanup the application"""
    # Startup
    print("Starting DocuChat application...")
    
    # Initialize OpenAI components after environment variables are loaded
    global embeddings, llm
    
    # Check if API key is available
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("❌ OPENAI_API_KEY not found in environment variables")
        print("Make sure you have a .env file with OPENAI_API_KEY=your-key")
        return
    else:
        print(f"✅ Found OpenAI API key: {api_key[:10]}...{api_key[-4:]}")
    
    try:
        print("🔄 Initializing OpenAI embeddings...")
        # Set environment variable first
        os.environ["OPENAI_API_KEY"] = api_key
        
        # Try alternative initialization approaches
        try:
            embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
        except Exception as e1:
            print(f"⚠️ First attempt failed: {e1}")
            try:
                # Fallback with minimal parameters
                from langchain_openai.embeddings import OpenAIEmbeddings as AltOpenAIEmbeddings
                embeddings = AltOpenAIEmbeddings()
            except Exception as e2:
                print(f"⚠️ Second attempt failed: {e2}")
                # Final fallback - use older style initialization
                from langchain.embeddings import OpenAIEmbeddings as LegacyOpenAIEmbeddings
                embeddings = LegacyOpenAIEmbeddings()
        
        print("✅ OpenAI embeddings initialized successfully")
        
        print("🔄 Initializing OpenAI chat model...")
        llm = ChatOpenAI(temperature=0.7, model="gpt-3.5-turbo")
        print("✅ OpenAI chat model initialized successfully")
        
    except Exception as e:
        print(f"❌ Error initializing OpenAI components: {e}")
        print(f"❌ Error type: {type(e).__name__}")
        import traceback
        print(f"❌ Full traceback: {traceback.format_exc()}")
        print("Make sure your OPENAI_API_KEY is correct and has sufficient credits")
    
    # Check ChromaDB connection
    try:
        # Try to list collections to test connection
        collections = chroma_client.list_collections()
        print(f"✅ ChromaDB connected successfully. Found {len(collections)} collections.")
    except Exception as e:
        print(f"⚠️ ChromaDB connection issue: {e}")
        print("Continuing with in-memory vector store as fallback")
    
    yield
    
    # Shutdown (cleanup if needed)
    print("Shutting down DocuChat application...")

# Initialize FastAPI app with lifespan
app = FastAPI(
    title="DocuChat", 
    description="Simple Document Q&A System",
    lifespan=lifespan
)

@app.get("/")
async def root():
    """Root endpoint - returns simple HTML interface"""
    html_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>DocuChat - Document Q&A</title>
        <style>
            body { font-family: Arial, sans-serif; margin: 40px; }
            .container { max-width: 800px; margin: 0 auto; }
            .upload-section, .chat-section { margin: 20px 0; padding: 20px; border: 1px solid #ddd; }
            input, textarea, button { margin: 10px 0; padding: 10px; }
            textarea { width: 100%; height: 100px; }
            button { background: #007bff; color: white; border: none; cursor: pointer; }
            .response { margin: 10px 0; padding: 10px; background: #f8f9fa; }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>📚 DocuChat - Document Q&A System</h1>
            
            <div class="upload-section">
                <h3>1. Upload Documents</h3>
                <input type="file" id="fileInput" multiple accept=".pdf,.txt,.docx">
                <button onclick="uploadFiles()">Upload Documents</button>
                <div id="uploadStatus"></div>
            </div>
            
            <div class="chat-section">
                <h3>2. Ask Questions</h3>
                <textarea id="questionInput" placeholder="Ask a question about your documents..."></textarea>
                <button onclick="askQuestion()">Ask Question</button>
                <div id="chatResponse"></div>
            </div>
        </div>

        <script>
            async function uploadFiles() {
                const fileInput = document.getElementById('fileInput');
                const files = fileInput.files;
                const status = document.getElementById('uploadStatus');
                
                if (files.length === 0) {
                    status.innerHTML = '<p style="color: red;">Please select files first!</p>';
                    return;
                }
                
                status.innerHTML = '<p>Uploading...</p>';
                
                for (let file of files) {
                    const formData = new FormData();
                    formData.append('file', file);
                    
                    try {
                        const response = await fetch('/upload', {
                            method: 'POST',
                            body: formData
                        });
                        
                        if (response.ok) {
                            status.innerHTML += `<p style="color: green;">✓ ${file.name} uploaded successfully!</p>`;
                        } else {
                            status.innerHTML += `<p style="color: red;">✗ Failed to upload ${file.name}</p>`;
                        }
                    } catch (error) {
                        status.innerHTML += `<p style="color: red;">✗ Error uploading ${file.name}</p>`;
                    }
                }
            }
            
            async function askQuestion() {
                const questionInput = document.getElementById('questionInput');
                const question = questionInput.value.trim();
                const responseDiv = document.getElementById('chatResponse');
                
                if (!question) {
                    responseDiv.innerHTML = '<p style="color: red;">Please enter a question!</p>';
                    return;
                }
                
                responseDiv.innerHTML = '<p>Thinking...</p>';
                
                try {
                    const response = await fetch('/chat', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({ question: question })
                    });
                    
                    const data = await response.json();
                    
                    if (response.ok) {
                        responseDiv.innerHTML = `
                            <div class="response">
                                <h4>Question: ${question}</h4>
                                <p><strong>Answer:</strong> ${data.answer}</p>
                                ${data.sources ? `<p><strong>Sources:</strong> ${data.sources.join(', ')}</p>` : ''}
                            </div>
                        `;
                    } else {
                        responseDiv.innerHTML = `<p style="color: red;">Error: ${data.detail}</p>`;
                    }
                } catch (error) {
                    responseDiv.innerHTML = '<p style="color: red;">Error connecting to server!</p>';
                }
                
                questionInput.value = '';
            }
        </script>
    </body>
    </html>
    """
    return HTMLResponse(content=html_content)

@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload and process a document for Q&A"""
    
    print(f"🔄 Starting upload process for: {file.filename}")
    
    # Check if OpenAI components are initialized
    if embeddings is None:
        print("❌ OpenAI embeddings not initialized")
        raise HTTPException(status_code=500, detail="OpenAI embeddings not initialized. Check your API key.")
    
    # Check if file type is supported
    supported_types = ['.pdf', '.txt', '.docx']
    file_extension = os.path.splitext(file.filename)[1].lower()
    
    if file_extension not in supported_types:
        print(f"❌ Unsupported file type: {file_extension}")
        raise HTTPException(status_code=400, detail=f"Unsupported file type. Supported: {supported_types}")
    
    print(f"✅ File type {file_extension} is supported")
    
    tmp_file_path = None
    
    try:
        print("📁 Saving uploaded file temporarily...")
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
            content = await file.read()
            tmp_file.write(content)
            tmp_file_path = tmp_file.name
        
        print(f"✅ File saved to: {tmp_file_path}")
        
        print("📄 Extracting text from file...")
        # Extract text from the file based on its type
        documents = extract_text_from_file(tmp_file_path, file.filename, file_extension)
        print(f"✅ Extracted {len(documents)} document(s)")
        
        if not documents or not any(doc.page_content.strip() for doc in documents):
            print("❌ No text content found in document")
            raise Exception("No text content could be extracted from the file")
        
        print("✂️ Splitting documents into chunks...")
        # Split documents into chunks for better processing
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        chunks = text_splitter.split_documents(documents)
        print(f"✅ Created {len(chunks)} chunks")
        
        if not chunks:
            print("❌ No chunks created from documents")
            raise Exception("Could not create text chunks from the document")
        
        print("🔗 Adding chunks to vector store...")
        # Add chunks to vector store
        global vectorstore, qa_chain
        try:
            if vectorstore is None:
                print("🆕 Creating new vector store...")
                # Create new vector store with simple in-memory approach first
                vectorstore = Chroma.from_documents(chunks, embeddings)
                print("✅ Created new in-memory vector store")
            else:
                print("➕ Adding to existing vector store...")
                vectorstore.add_documents(chunks)
                print("✅ Added documents to existing vector store")
        except Exception as ve:
            print(f"❌ Vector store error: {ve}")
            print("🔄 Trying simple in-memory vector store...")
            # Fallback: create simple in-memory vector store
            vectorstore = Chroma.from_documents(chunks, embeddings)
            print("✅ Created fallback vector store")
        
        print("🤖 Creating Q&A chain...")
        # Create/update the Q&A chain
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),
            return_source_documents=True
        )
        print("✅ Q&A chain created successfully")
        
        # Clean up temporary file
        print("🧹 Cleaning up temporary file...")
        os.unlink(tmp_file_path)
        print("✅ Cleanup completed")
        
        print(f"🎉 Successfully processed {file.filename} with {len(chunks)} chunks")
        return {"message": f"Successfully processed {file.filename}", "chunks": len(chunks)}
        
    except Exception as e:
        print(f"❌ Error during processing: {str(e)}")
        print(f"❌ Error type: {type(e).__name__}")
        import traceback
        print(f"❌ Full traceback: {traceback.format_exc()}")
        
        # Clean up temporary file if it exists
        if tmp_file_path and os.path.exists(tmp_file_path):
            try:
                os.unlink(tmp_file_path)
                print("🧹 Cleaned up temporary file after error")
            except Exception as cleanup_error:
                print(f"⚠️ Could not clean up temp file: {cleanup_error}")
        
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

def extract_text_from_file(file_path: str, filename: str, file_extension: str) -> List[Document]:
    """Extract text from different file types and return as LangChain documents"""
    
    print(f"📖 Extracting text from {file_extension} file: {filename}")
    documents = []
    
    try:
        if file_extension == '.pdf':
            print("📕 Processing PDF file...")
            # Handle PDF files
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                print(f"📄 PDF has {len(pdf_reader.pages)} pages")
                text = ""
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    print(f"✅ Extracted text from page {i+1}: {len(page_text)} characters")
                
                if text.strip():
                    documents.append(Document(page_content=text, metadata={"source": filename}))
                    print(f"✅ Created document with {len(text)} characters")
                else:
                    print("❌ No text extracted from PDF")
                
        elif file_extension == '.txt':
            print("📄 Processing text file...")
            # Handle text files
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                print(f"✅ Read {len(text)} characters from text file")
                if text.strip():
                    documents.append(Document(page_content=text, metadata={"source": filename}))
                else:
                    print("❌ Text file is empty")
                
        elif file_extension == '.docx':
            print("📘 Processing Word document...")
            # Handle Word documents
            doc = docx.Document(file_path)
            print(f"📄 Document has {len(doc.paragraphs)} paragraphs")
            text = ""
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text
                text += para_text + "\n"
                if para_text.strip():
                    print(f"✅ Paragraph {i+1}: {len(para_text)} characters")
            
            if text.strip():
                documents.append(Document(page_content=text, metadata={"source": filename}))
                print(f"✅ Created document with {len(text)} characters")
            else:
                print("❌ No text extracted from Word document")
            
        print(f"📋 Total documents created: {len(documents)}")
        return documents
            
    except Exception as e:
        print(f"❌ Error extracting text from {file_extension}: {str(e)}")
        import traceback
        print(f"❌ Extraction traceback: {traceback.format_exc()}")
        raise Exception(f"Error reading {file_extension} file: {str(e)}")

@app.post("/chat", response_model=ChatResponse)
async def chat_with_documents(request: ChatRequest):
    """Ask a question about the uploaded documents"""
    
    # Check if OpenAI components are initialized
    if llm is None or embeddings is None:
        raise HTTPException(status_code=500, detail="OpenAI components not initialized. Check your API key.")
    
    # Check if we have documents loaded
    if qa_chain is None:
        raise HTTPException(status_code=400, detail="No documents uploaded yet. Please upload documents first.")
    
    try:
        # Get answer from the Q&A chain
        result = qa_chain({"query": request.question})
        
        # Extract source information
        sources = []
        if "source_documents" in result:
            sources = [doc.metadata.get("source", "Unknown") for doc in result["source_documents"]]
            # Remove duplicates while preserving order
            sources = list(dict.fromkeys(sources))
        
        return ChatResponse(
            answer=result["result"],
            sources=sources
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing question: {str(e)}")

@app.get("/health")
async def health_check():
    """Simple health check endpoint"""
    return {
        "status": "healthy", 
        "message": "DocuChat is running!",
        "openai_embeddings_ready": embeddings is not None,
        "openai_llm_ready": llm is not None,
        "vectorstore_ready": vectorstore is not None,
        "qa_chain_ready": qa_chain is not None
    }

if __name__ == "__main__":
    import uvicorn
    # Check if OpenAI API key is loaded from .env file
    if not os.getenv("OPENAI_API_KEY"):
        print("Warning: OPENAI_API_KEY not found!")
        print("Please make sure you have a .env file with: OPENAI_API_KEY=your-api-key-here")
    else:
        print("✅ OpenAI API key loaded successfully from .env file")
    
    uvicorn.run(app, host="0.0.0.0", port=8000)
